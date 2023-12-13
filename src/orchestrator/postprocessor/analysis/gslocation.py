#! /usr/bin/env python3

# Copyright (C)
# Author: alberto-ferrero

import os
import time
import geopandas as gpd
import matplotlib.pyplot as plt

#Import docx
from docx import Document

#HACK ignoring all conversion and deprecations WARNINGS
import warnings
warnings.simplefilter(action='ignore')

""" E2E Performance Simulator Analysis: Ground Stations locations """

def write(doc: Document, outputPlotFolderPath: str, simulationRequest: dict):
    """Write analyis chapter """
    tick = time.time()
    #Read from Simulation Request, the input ground-stations
    gss = []
    for gs in simulationRequest['groundstations']:
        gss.append({
            "id": gs['id'],
            "latitude": gs['location']['latitude'],
            "longitude": gs['location']['longitude']
        })
    
    if len(gss) > 0:

        #Build image
        worldmap = gpd.read_file(gpd.datasets.get_path("naturalearth_lowres"))
        fig, ax = plt.subplots(figsize=(8, 8))
        worldmap.plot(color="lightgrey", ax=ax)
        ax.set(xlim=[-180, 180], ylim=[-90, 90])
        for gs in gss:
          lat = gs['latitude']
          lng = gs['longitude']
          ax.scatter(lng, lat, s=40, c=['red'], alpha=0.9)
          ax.annotate(gs['id'], (lng, lat))
        ax.set_xlabel("Longitude [deg]")
        ax.set_ylabel("Latitude [deg]")
        figPath = os.path.join(outputPlotFolderPath, "analysis_groundstations-location.png")
        fig.tight_layout()
        fig.savefig(figPath,
                    bbox_inches='tight')

        #Write       
        doc.add_heading("Ground Stations Locations", 1)
        doc.add_paragraph('The picture below shows ground stations locations, as defined in configuration')
        p = doc.add_paragraph()
        r = p.add_run()
        r.add_picture(figPath)
        
        doc.add_page_break()

        print('   - Added section on Ground Stations Locations in {:.4f} seconds'.format(time.time() - tick))


# -*- coding: utf-8 -*-