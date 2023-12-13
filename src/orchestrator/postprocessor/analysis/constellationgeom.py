#! /usr/bin/env python3

# Copyright (C)
# Author: alberto-ferrero

import os
import glob
import time
import pyorb
import geopandas as gpd
import matplotlib.pyplot as plt
import pandas as pd
import numpy as np
from PIL import Image
from math import pi

from ..analysis.noc import getSatelliteLatitudeLongitude, getSatellitePositionVelocity
from ....utils.filemanager import makeOutputFolder
from ....utils.timeconverter import getTimestampFromDate

#HACK ignoring all conversion and deprecations WARNINGS
import warnings
warnings.simplefilter(action='ignore')

#Import docx
from docx import Document

""" E2E Performance Simulator Analysis: Constellation Geometry """

def write(doc: Document, simulationRequest: dict, outputPlotFolderPath: str, flightDynamicsDataOutputPath: str):
    """Write analyis chapter """
    #Read from Flight Dynamics file and extract mean Keplerian elements
    tick = time.time()
    constDf = {}
    utcTimestamps = []
    for fileName in glob.glob(os.path.join(flightDynamicsDataOutputPath, "*_orbit-state.csv")):
        df = pd.read_csv(fileName)
        satId = fileName.split(os.sep)[-1].split("_")[0]
        df = df[['utcTime', 'X', 'Y', 'Z', 'Vx', 'Vy', 'Vz']]
        df['timestamp'] = df['utcTime'].apply(getTimestampFromDate)
        constDf[satId] = df
        #Get list of timestamps
        if utcTimestamps == []:
            utcTimestamps = df['timestamp'].to_list()
        else:
            utcTimestamps = set(utcTimestamps).intersection(df['timestamp'].to_list())

    if len(constDf) == 0:
        return

    #Get Keplerian Elements from orbit definition
    raan = []
    la = []
    for satId in constDf:
        kepl = pyorb.cart_to_kep(getSatellitePositionVelocity(constDf[satId], index=0)).tolist()
        raan.append(kepl[4] * 180.0 / pi)
        la.append((kepl[3] + kepl[5]) * 180.0 / pi)
        la[-1] = la[-1] if la[-1] < 359.99 else la[-1] - 359.99

    #Get list of timestamps
    [set(utcTimestamps)].sort()

    #Save image for each time in tmp
    tmpPath = os.path.join(outputPlotFolderPath, 'tmp')
    makeOutputFolder(tmpPath)
    images = []
    for i, utcTimestamp in enumerate(utcTimestamps):
        if i > 100:
            break
        #Build map
        worldmap = gpd.read_file(gpd.datasets.get_path("naturalearth_lowres"))
        fig, ax = plt.subplots(figsize=(8, 8))
        ax.set_xlabel("Longitude [deg]")
        ax.set_ylabel("Latitude [deg]")
        worldmap.plot(color="darkgrey", ax=ax)
        ax.set(xlim=[-180, 180], ylim=[-90, 90])
        for satId in constDf:
            index = np.where(constDf[satId]["timestamp"] == utcTimestamp)[0].tolist()[0]
            lat, lng = getSatelliteLatitudeLongitude(constDf[satId], index=index)
            ax.scatter(lng, lat, s=10, c=['black'], alpha=0.7)
            ax.annotate(satId, (lng, lat), fontsize=7)
        t = constDf[satId].iloc[index]['utcTime']
        ax.set_title(t)
        ax.margins(x=0.9,y=0)
        figPath = os.path.join(tmpPath, "analysis_constellation-orbit-{}.jpg".format(constDf[satId].iloc[index]['timestamp']))
        images.append(figPath)
        fig.savefig(figPath)
        #Save first, image and position
        if i == 0:
            timezero = t.replace("T", " at ").replace("Z", "")
            figPath = os.path.join(outputPlotFolderPath, "analysis_constellation-orbit.jpg")
            fig.savefig(figPath, bbox_inches='tight')

    #Save gifs
    images.sort()
    frames = [Image.open(image) for image in images]
    frame = frames[0]
    frame.save(os.path.join(outputPlotFolderPath, 'analysis_constellation-orbit.gif'), 
                   format="GIF", append_images=frames,
                   save_all=True, duration=200, loop=0)

    #Write
    doc.add_heading("Constellation Geometry", 1)

    doc.add_paragraph('The picture below shows the constellation orbit print, at simulation time zero: {}'.format(timezero))
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_picture(figPath)

    #Save geometry
    fig, ax = plt.subplots(figsize=(8, 8))
    ax.set_xlabel("RAAN [deg]")
    ax.set_ylabel("Latitude Argument [deg]")
    ax.set(xlim=[-1, 166], ylim=[-1, 360])
    ax.scatter(raan, la, s=10, c=['black'], alpha=0.7)
    figPath = os.path.join(outputPlotFolderPath, "analysis_constellation-geometry.jpg")
    fig.tight_layout()
    fig.savefig(figPath, bbox_inches='tight')
   
    doc.add_paragraph('The picture below shows the constellation topology, at simulation time zero: {}'.format(timezero))
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_picture(figPath)
    
    doc.add_page_break()
    
    print('   - Added section on Constellation Geometry in {:.4f} seconds'.format(time.time() - tick))

# -*- coding: utf-8 -*-