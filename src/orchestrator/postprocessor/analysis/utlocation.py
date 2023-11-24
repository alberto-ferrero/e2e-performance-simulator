#! /usr/bin/env python3

# Copyright (C)
# Author: alberto-ferrero

import os
import geopandas as gpd
import matplotlib.pyplot as plt
import time

#Import docx
from docx import Document

#HACK ignoring all conversion and deprecations WARNINGS
import warnings
warnings.simplefilter(action='ignore')

""" E2E Performance Simulator Analysis: User Terminals locations """

def write(doc: Document, outputPlotFolderPath: str, simulationRequest: dict):
    """Write analyis chapter """
    tick = time.time()
    #Read from Simulation Request, the input user-terminals
    gss = []
    for gs in simulationRequest['userterminals']:
        gss.append({
            "id": gs['id'],
            "latitude": gs['location']['latitude'],
            "longitude": gs['location']['longitude']
        })
    
    if len(gss) > 0:

        #Build image
        worldmap = gpd.read_file(gpd.datasets.get_path("naturalearth_lowres"))
        fig, ax = plt.subplots()
        worldmap.plot(color="lightgrey", ax=ax)
        for gs in gss:
          lat = gs['latitude']
          lng = gs['longitude']
          ax.scatter(lng, lat, s=40, c=['red'], alpha=0.9)
          ax.annotate(gs['id'], (lng, lat))
        ax.set_xlabel("Longitude [deg]")
        ax.set_ylabel("Latitude [deg]")
        figPath = os.path.join(outputPlotFolderPath, "analysis_userterminals-location.png")
        fig.savefig(figPath,
                    bbox_inches='tight')

        #Write       
        doc.add_heading("User Terminals Locations", 1)
        doc.add_paragraph('The picture below shows user terminals locations, as defined in configuration')
        p = doc.add_paragraph()
        r = p.add_run()
        r.add_picture(figPath)
        
        doc.add_page_break()
    
        print(' - Added section on User Terminals Locations in  {:.4f} seconds'.format(time.time() - tick))


# -*- coding: utf-8 -*-