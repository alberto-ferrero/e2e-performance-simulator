#! /usr/bin/env python3

# Copyright (C)
# Author: alberto-ferrero

import os
import glob
import time

import geopandas as gpd
import matplotlib.pyplot as plt
from PIL import Image

#HACK ignoring all conversion and deprecations WARNINGS
import warnings
warnings.simplefilter(action='ignore')

import pandas as pd
import numpy as np

from ..analysis.noc import getFlatXConnections, getItalXConnections
from ..analysis.noc import getCloserSatelliteDistance, getCloserSatelliteDistanceMesh
from ..analysis.noc import getPointFromLatLong, getSatelliteLatitudeLongitude, getDistance, getSatellitePosition

from ....utils.filemanager import makeOutputFolder
from ....utils.timeconverter import getDatetimeFromDate

#Import docx
from docx import Document

#Constants
Re = 6371000       #[m]
c = 299792458      #[m/s]
deltaAngle = 5     #[deg]

""" E2E Performance Simulator Analysis: Latency """

def write(doc: Document, outputDataFolderPath: str, outputPlotFolderPath: str, flightDynamicsDataOutputPath: str):
    tick = time.time()
    #Read from Flight Dynamics file and extract EME2000 coordinates
    satsDf = {}
    totSats = 0
    totPlanes = 0
    for fileName in glob.glob(os.path.join(flightDynamicsDataOutputPath, "*_orbit-state.csv")):
        satId = fileName.split(os.sep)[-1].split("_")[0]
        df = pd.read_csv(fileName)
        df = df[['utcTime', 'X', 'Y', 'Z']].iloc[0:1]
        satsDf[satId] = df
        t = getDatetimeFromDate(df.iloc[0]['utcTime'])
        totSats = int(satId.split("-")[-1]) if int(satId.split("-")[-1]) > totSats else totSats
        totPlanes = int(satId.split("-")[-2].replace('P','')) if int(satId.split("-")[-2].replace('P','')) > totPlanes else totPlanes 

    if satsDf == {}:
        return
    
    outputAnalysFolderPath = os.path.join(outputDataFolderPath, 'analysis')
    makeOutputFolder(outputAnalysFolderPath)

    tmpPath = os.path.join(outputPlotFolderPath, 'tmp')
    makeOutputFolder(tmpPath)

    #Write section       
    doc.add_heading("Signal Latency", 1)

    #Iterate considering both meshes
    lats = np.arange(-90, 90 + deltaAngle / 2.0, deltaAngle)
    lngs = np.arange(-180, 180 + deltaAngle / 2.0, deltaAngle)
    for getMesh, tag in ((getFlatXConnections, 'Flat X'), (getItalXConnections, 'Ital X')):
        #Move around globe from focal point and calculate transmission based on speed of light
        latitudes = np.arange(0, 95, 5)
        #Save mesh
        with open(os.path.join(outputAnalysFolderPath, "analysis_connections-{}-mesh-time-0.txt".format(tag.replace(" ", ""))), "w") as f:
            for satId in satsDf:
                f.write("{} {}\n".format(satId, getMesh(satId, totPlanes, totSats)))
        images = []
        for latitude in latitudes:
            longitude = 0
            firstGeoPoint = getPointFromLatLong(latitude, longitude, t)

            #Get first communication distance
            firstDistance, firstSatId = getCloserSatelliteDistance(satsDf, firstGeoPoint)
            firstDelay = firstDistance / c

            #Empty delays
            delays = np.empty(shape=(len(lats), len(lngs)))
            worstDelay = 0
            worstLatLng = []
            #Get closer satellite distance and calculate delay as d * c
            for i, lat in enumerate(lats):
                for j, lng in enumerate(lngs):
                    geoPoint = getPointFromLatLong(lat, lng, t)
                    distance, closerSatId = getCloserSatelliteDistance(satsDf, geoPoint)
                    #Go through mesh from initial geo point up to target point
                    nextSatId = closerSatId
                    contactedSatIds = [closerSatId, ]
                    #[DEBUG]print("\n", lat, lng, "from SAT", closerSatId, "to SAT", firstSatId)
                    while nextSatId != firstSatId:
                        #From the current point, amongh the ones in the mesh, get closer to initial geo point
                        closerSatId = getCloserSatelliteDistanceMesh(getMesh, totPlanes, totSats, satsDf, firstGeoPoint, contactedSatIds + [nextSatId,], firstSatId)
                        #[DEBUG]print(nextSatId, [satId for satId in getMesh(nextSatId) if satId not in contactedSatIds], 'closer', closerSatId)
                        contactedSatIds.append(closerSatId)
                        #Add intersatellite distance and go to next
                        distance += getDistance(getSatellitePosition(satsDf[closerSatId]),
                                                getSatellitePosition(satsDf[nextSatId]))
                        nextSatId = closerSatId
                    #[DEBUG]if (distance / c + firstDelay) * 1000.0 > 300:
                        #[DEBUG]print('Delay:', (distance / c + firstDelay) * 1000.0)
                        #[DEBUG]print(lat, lng, "from SAT", contactedSatIds[0], "to SAT", firstSatId, "\n")
                        #[DEBUG]print(contactedSatIds + [nextSatId,])
                    delay = (distance / c + firstDelay) * 1000.0 #to [millis]
                    delays[i][j] = delay
                    if delay > worstDelay:
                        worstDelay = delay
                        worstLatLng = [lat, lng]

            #Build heat map
            worldmap = gpd.read_file(gpd.datasets.get_path("naturalearth_lowres"))
            fig, ax = plt.subplots(figsize=(8, 8))
            worldmap.plot(color="darkgrey", ax=ax)
            cs = ax.contourf(lngs, lats, delays, alpha=0.3)
            cf = fig.colorbar(cs, fraction=0.046, pad=0.04)
            cf.ax.set_ylabel("Latency [millis]", loc='center')
            ax.set_xlabel("Longitude [deg]")
            ax.set_ylabel("Latitude [deg]")
            ax.set(xlim=[lngs[0], lngs[-1]], ylim=[lats[0], lats[-1]])
            for satId in satsDf:
                lat, lng = getSatelliteLatitudeLongitude(satsDf[satId])
                ax.scatter(lng, lat, s=10, c=['black'], alpha=0.7)
                #[DEBUG]ax.annotate(satId, (lng, lat), fontsize=7)
            ax.scatter(longitude, latitude, s=30, c=['black'], alpha=0.9)
            ax.annotate("UT", (longitude, latitude))
            ax.scatter(worstLatLng[1], worstLatLng[0], s=30, c=['red'], alpha=0.9, marker="*")
            ax.annotate("WUL", (worstLatLng[1], worstLatLng[0]))
            #Save initial figure
            if latitude == 0:
                figPath = os.path.join(outputPlotFolderPath, "analysis_latency-{}-mesh.jpg".format(tag.replace(" ", "")))
                fig.tight_layout()
                fig.savefig(figPath, bbox_inches='tight')
                doc.add_paragraph('The picture below shows signal delay in milliseconds, from an user terminal set at Latitude {} deg, Longitude {} deg, for {} mesh geometry'.format(latitude, longitude, tag))
                p = doc.add_paragraph()
                r = p.add_run()
                r.add_picture(figPath)
                doc.add_paragraph('The table below summarizes position of WUL, Worst User Location, considering transmission delay from UT at different latitudes')
                table = doc.add_table(rows=1, cols=2, style="Table Grid")
                heading = table.rows[0].cells
                heading[0].text = "UT Coordinates"
                heading[1].text = "WUL Coordinates"
            ax.set_title("Lat = {} deg".format(latitude))
            pd.DataFrame(delays, columns=lngs, index=lats).to_csv(os.path.join(outputAnalysFolderPath, 'analysis_latency-{}-mesh-{}.csv'.format(tag.replace(" ", ""), latitude)), index=False)
            # Set value of worst condition
            cells = table.add_row().cells
            cells[0].text = "Lng = 00 deg\nLat = {} deg".format(str(int(latitude)).zfill(2))
            cells[1].text = "Lng = {} deg\nLat = {} deg\nDelay = {} ms".format(str(int(worstLatLng[1])).zfill(2), str(int(worstLatLng[0])).zfill(2), str(int(worstDelay)).zfill(3))

            figTmpPath = os.path.join(tmpPath, "analysis_latency-{}-mesh-{}.jpg".format(tag.replace(" ", ""), latitude))
            fig.tight_layout()
            fig.savefig(figTmpPath, bbox_inches='tight')
            images.append(figTmpPath)

        #Save gif
        frames = [Image.open(image) for image in images]
        frame = frames[0]
        frame.save(os.path.join(outputPlotFolderPath, 'analysis_latency-{}-mesh.gif'.format(tag)), 
                    format="GIF", append_images=frames,
                    save_all=True, duration=100, loop=0)
        plt.close('all')
    
    doc.add_page_break()
    print('   - Added section on User Signal Latency in {:.4f} seconds'.format(time.time() - tick))

# -*- coding: utf-8 -*-