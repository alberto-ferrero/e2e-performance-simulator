#! /usr/bin/env python3

# Copyright (C)
# Author: alberto-ferrero

import os
import glob
import time

import pandas as pd
import numpy as np

import geopandas as gpd
import matplotlib.pyplot as plt
from PIL import Image

#HACK ignoring all conversion and deprecations WARNINGS
import warnings
warnings.simplefilter(action='ignore')

from ..analysis.noc import getCloserSatelliteDistance, getCloserSatelliteContactDistance
from ..analysis.noc import getGeopointFromLatLong, getSatelliteLatitudeLongitude

from ....utils.filemanager import makeOutputFolder
from ....utils.timeconverter import getDateFromTimestamp, getTimestampFromDate

from ....spacelink.request import getContactStates

#Import docx
from docx import Document

#Constants
Re = 6371000       #[m]
c = 299792458      #[m/s]
deltaAngle = 5     #[deg]

""" E2E Performance Simulator Analysis: Useer Terminals network connection """

def write(doc: Document, outputDataFolderPath: str, outputPlotFolderPath: str, flightDynamicsDataOutputPath: str):
    tick = time.time()    
    
    #Extract intersatellite visibility
    utcTimestamps = []
    satsStatesDf = {}
    satsContactsDf = {}
    for fileName in glob.glob(os.path.join(flightDynamicsDataOutputPath, "*_orbit-state.csv")):
        satId = fileName.split(os.sep)[-1].split("_")[0]
        df = pd.read_csv(fileName)
        df = df[['utcTime', 'X', 'Y', 'Z']]
        df['timestamp'] = df['utcTime'].apply(getTimestampFromDate)
        satsStatesDf[satId] = df
        #Get list of timestamps
        if utcTimestamps == []:
            utcTimestamps = df['timestamp'].to_list()
        else:
            utcTimestamps = set(utcTimestamps).intersection(df['timestamp'].to_list())
        #Read contacts
        df = pd.read_csv(os.path.join(flightDynamicsDataOutputPath, satId + "_contacts.csv"))
        df['startTimestamp'] = df['startUtcTime'].apply(getTimestampFromDate)
        df['endTimestamp'] = df['endUtcTime'].apply(getTimestampFromDate)
        df = df[df['contactType'] == 'ISV']
        satsContactsDf[satId] = df

    if satsStatesDf == {}:
        return
    
    #Get list of timestamps
    [set(utcTimestamps)].sort()

    outputAnalysFolderPath = os.path.join(outputDataFolderPath, 'analysis')
    makeOutputFolder(outputAnalysFolderPath)

    tmpPath = os.path.join(outputPlotFolderPath, 'tmp')
    makeOutputFolder(tmpPath)

    #Write section
    doc.add_heading("Netowrk Connections", 1)

    #Set up UT locations
    rivadaSpaceNetworks = (48.144802679076044, 11.595048249556703)
    terranOrbital = (33.683465115437244, -117.77589876915086)
    images = []
    for i, utcTimestamp in enumerate(utcTimestamps):
        if i > 100:
            break
        #Build map
        worldmap = gpd.read_file(gpd.datasets.get_path("naturalearth_lowres"))
        fig, ax = plt.subplots(figsize=(8, 8))
        ax.set_xlabel("Longitude [deg]")
        ax.set_ylabel("Latitude [deg]")
        worldmap.plot(color="darkgrey", ax=ax)
        ax.set(xlim=[-180, 180], ylim=[-90, 90])
        t = getDateFromTimestamp(utcTimestamp)
        #Plot ground
        ax.scatter(terranOrbital[1], terranOrbital[0], s=30, c=['red'], alpha=0.9, marker="*")
        ax.annotate("TO", (terranOrbital[1], terranOrbital[0]))
        ax.scatter(rivadaSpaceNetworks[1], rivadaSpaceNetworks[0], s=30, c=['red'], alpha=0.9, marker="*")
        ax.annotate("RSN", (rivadaSpaceNetworks[1], rivadaSpaceNetworks[0]))
        geoTerran = getGeopointFromLatLong(terranOrbital[0], terranOrbital[1], t)
        geoRsn = getGeopointFromLatLong(rivadaSpaceNetworks[0], rivadaSpaceNetworks[1], t)
        #Plot satellites
        for satId in satsStatesDf:
            index = np.where(satsStatesDf[satId]["timestamp"] == utcTimestamp)[0].tolist()[0]
            lat, lng = getSatelliteLatitudeLongitude(satsStatesDf[satId], index=index)
            ax.scatter(lng, lat, s=10, c=['black'], alpha=0.7)
            ax.annotate(satId, (lng, lat), fontsize=6)
        #Get mesh to communicate between two points (TO -> RSN)
        _, firstSatId = getCloserSatelliteDistance(satsStatesDf, geoTerran, utcTimestamp)
        _, lastSatId = getCloserSatelliteDistance(satsStatesDf, geoRsn, utcTimestamp)
        #Go through mesh from initial geo point up to target point
        nextSatId = lastSatId
        contactedSatIds = [lastSatId, ]
        #[DEBUG]print("closest to TO", firstSatId)
        while nextSatId != firstSatId:
            #From the current point, amongh the ones in the mesh, get closer
            _, closerSatId = getCloserSatelliteContactDistance(utcTimestamp, nextSatId, geoTerran, satsStatesDf, satsContactsDf, contactedSatIds)
            if closerSatId == "":
                break
            contactedSatIds.append(closerSatId)
            nextSatId = closerSatId
        #Plot lines
        satId = contactedSatIds[0]
        index = np.where(satsStatesDf[satId]["timestamp"] == utcTimestamp)[0].tolist()[0]
        satLat, satLng = getSatelliteLatitudeLongitude(satsStatesDf[satId], index=index)
        for id in range(len(contactedSatIds)):
            if id == 0:
                ax.plot([rivadaSpaceNetworks[1], satLng], [rivadaSpaceNetworks[0], satLat], 'b-')
                if len(contactedSatIds) == 2:
                    satId = contactedSatIds[id+1]
                    index = np.where(satsStatesDf[satId]["timestamp"] == utcTimestamp)[0].tolist()[0]
                    nextSatLat, nextSatLng = getSatelliteLatitudeLongitude(satsStatesDf[satId], index=index)
                    ax.plot([nextSatLng, satLng], [nextSatLat, satLat], 'b-')
                    satLng = nextSatLng
                    satLat = nextSatLat
            if id < len(contactedSatIds) - 1:
                satId = contactedSatIds[id+1]
                index = np.where(satsStatesDf[satId]["timestamp"] == utcTimestamp)[0].tolist()[0]
                nextSatLat, nextSatLng = getSatelliteLatitudeLongitude(satsStatesDf[satId], index=index)
                ax.plot([nextSatLng, satLng], [nextSatLat, satLat], 'b-')
                satLng = nextSatLng
                satLat = nextSatLat
            else:
                ax.plot([terranOrbital[1], satLng], [terranOrbital[0], satLat], 'b-')
            #Save fig
            ax.set_title(t)
            ax.margins(x=0.9,y=0)
            figPath = os.path.join(tmpPath, "analysis_constellation-network-mesh-{}.jpg".format(utcTimestamp))
            images.append(figPath)
            fig.savefig(figPath)
        #Save first, image and position
        if i == 0:
            timezero = getDateFromTimestamp(utcTimestamp).replace("T", " at ").replace("Z", "")
            figPath = os.path.join(outputPlotFolderPath, "analysis_constellation-network-mesh.jpg")
            fig.savefig(figPath, bbox_inches='tight')                
            doc.add_paragraph('The picture below shows the constellation network communication, at simulation time zero: {}, between Rivada Space Network offices in Munich and Terran Orbital facility, in Irvine, for mesh'.format(timezero))
            p = doc.add_paragraph()
            r = p.add_run()
            r.add_picture(figPath)
    #Save gifs
    images.sort()
    frames = [Image.open(image) for image in images]
    if len(frames) > 0:
        frame = frames[0]
        frame.save(os.path.join(outputPlotFolderPath, 'analysis_constellation-network-mesh.gif'), 
                    format="GIF", append_images=frames,
                    save_all=True, duration=200, loop=0)
    plt.close('all')
    doc.add_page_break()
    print('   - Added section on User Terminals network topology in {:.4f} seconds'.format(time.time() - tick))

# -*- coding: utf-8 -*-