#! /usr/bin/env python3

# Copyright (C)
# Author: alberto-ferrero

import os
import glob
import time

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.dates as md

from math import ceil

from ....orchestrator.preprocessor.preprocessor import readGroundStations, readUserTerminals
from ....utils.filemanager import makeOutputFolder
from ....utils.timeconverter import getDatetimeFromDate

# Register time converters
pd.plotting.register_matplotlib_converters()

#Import docx
from docx import Document

""" E2E Performance Simulator Analysis: Satellite Visiblity per Latitude """

def write(doc: Document, simulationRequest: dict, outputDataFolderPath: str, outputPlotFolderPath: str, flightDynamicsDataOutputPath: str):
    """Write analyis chapter """
    from scipy.interpolate import RegularGridInterpolator as RGI
    tick = time.time()
    #Read from Flight Dynamics file and extract contacts info
    utsDf = {}
    gssDf = {}
    for fileName in glob.glob(os.path.join(flightDynamicsDataOutputPath, "*_contacts.csv")):
        satId = fileName.split(os.sep)[-1].split("_")[0]
        try:
            df = pd.read_csv(fileName)
            if df.empty:
                continue
        except:
            continue
        df = df.loc[df['contactType'] == "POI"]
        if len(df) == 0:
            continue
        df = df[['argumentOfInterestId', 'startUtcTime', 'endUtcTime', 'maxElevationElevation']]
        df['satId'] = satId
        # Store contacts for the args
        argIds = set(df['argumentOfInterestId'].to_list())
        for argId in argIds:
            filteredDf = df.copy()
            filteredDf = filteredDf[filteredDf['argumentOfInterestId'] == argId]
            if 'ut-' in argId:
                if argId not in utsDf:
                    utsDf[argId] = filteredDf
                else:
                    utsDf[argId] = pd.concat([utsDf[argId], filteredDf], ignore_index=True)
                utsDf[argId].index
            if 'gs-' in argId:
                if argId not in gssDf:
                    gssDf[argId] = filteredDf
                else:
                    gssDf[argId] = pd.concat([gssDf[argId], filteredDf], ignore_index=True)
                gssDf[argId].index

    if gssDf == {} and utsDf == {}:
        return

    outputAnalysFolderPath = os.path.join(outputDataFolderPath, 'analysis')
    makeOutputFolder(outputAnalysFolderPath)

    #Read coordinates of each GS and UT
    groundstations = readGroundStations(simulationRequest)
    userterminals = readUserTerminals(simulationRequest)
    coords = {}
    for gs in groundstations:
        if gs['id'] in gssDf.keys():
            coords[gs['id']] = {
                "lat": gs['location']['latitude'],
                "lng": gs['location']['longitude']
            }
    for ut in userterminals:
        if ut['id'] in utsDf.keys():
            coords[ut['id']] = {
                "lat": ut['location']['latitude'],
                "lng": ut['location']['longitude']
            }
    
    #Write       
    doc.add_heading("Propagated Contacts", 1)

    ###################################################################
    #Plot number of visible satellites for each User Terminal and Ground Station
    xfmt = md.DateFormatter('%Y-%m-%d %H')
    deltaAngle = 1
    lats = np.arange(0, 90 + deltaAngle / 2.0, deltaAngle)
    elevs = np.arange(0, 90 + deltaAngle / 2.0, deltaAngle)
    figPathV = ""
    totNumerosity = 0
    for t, d, descr in (("GS", gssDf, 'ground stations'), ("UT", utsDf, 'user terminals')):
        if len(d) == 0:
            continue
        ###################################################################
        #Set up figure
        fig, ax = plt.subplots(len(d), sharex=True, figsize=(7, 7))
        plt.subplots_adjust(hspace=.02)
        fig.suptitle("Satellite Visibility for {}s".format(t))
        #Set to list if only 1 subplot
        if len(d) == 1:
            ax = (ax,)

        #Add mapping for number of satellite visible and numerosity of the case
        visibility = np.empty(shape=(len(elevs), len(lats)))
        numerosity = np.zeros(shape=(len(elevs), len(lats)))
        
        maxContacts: list = []
        dates: list = []
        fileName = 'analysis_' + t + '-sat-visibility'
        for idx, (poiId, df) in enumerate(d.items()):
            
            ###############################################################
            #Extract number of contacts
            startTimeDf: pd.DataFrame = df[['satId', 'startUtcTime', 'maxElevationElevation']].copy()
            startTimeDf.rename(columns={"startUtcTime": "time"}, inplace=True)
            startTimeDf['contact'] = 1
            endTimeDf: pd.DataFrame = df[['satId', 'endUtcTime', 'maxElevationElevation']].copy()
            endTimeDf.rename(columns={"endUtcTime": "time"}, inplace=True)
            endTimeDf['contact'] = -1
            contactTimeDf = pd.concat([startTimeDf, endTimeDf], ignore_index=True)
            contactTimeDf['time'] = contactTimeDf['time'].apply(getDatetimeFromDate)
            contactTimeDf.sort_values(by='time', inplace=True)
            pd.to_datetime(contactTimeDf['time'], format='%Y-%m-%dT%H:%M:%S.%fZ')
            dates.extend([contactTimeDf['time'].to_list()[0], contactTimeDf['time'].to_list()[-1]])
            contactTimeDf['contacts'] = contactTimeDf['contact'].cumsum()
            maxContacts.append(max(contactTimeDf['contacts'].to_list()))
            #Plot
            ax[idx].plot(contactTimeDf['time'], contactTimeDf['contacts'], drawstyle="steps-post",  linewidth=0.3)

            #Add average heat map of satellite visibility, by latitude (specularty assumed)
            if poiId in coords:
                latIndex = int(ceil(abs(coords[poiId]['lat']) / 90 * len(lats)))
                latIndex = latIndex if latIndex < len(lats) else len(lats) - 1
                for elevIndex, elev in enumerate(elevs):
                    filteredElevDf = contactTimeDf[['contact', 'maxElevationElevation']].copy()
                    filteredElevDf = filteredElevDf[filteredElevDf['maxElevationElevation'] < elev]
                    if not filteredElevDf.empty:
                        filteredElevDf['filtContacts'] = filteredElevDf['contact'].cumsum()
                        visibility[elevIndex][latIndex] += filteredElevDf['filtContacts'].mean()
                        numerosity[elevIndex][latIndex] += 1
                    del filteredElevDf

        #Resize and set tags
        maxContacts = max(maxContacts)
        yTicks = np.arange(0, maxContacts + 1, int(maxContacts / 3)).tolist()
        minDate = min(dates)
        maxDate = max(dates)
        for idx, id in enumerate(d.keys()):
            ax[idx].text(0.97, 0.5, id, horizontalalignment='left', verticalalignment='center', transform=ax[idx].transAxes)
            ax[idx].xaxis.set_major_formatter(xfmt)
            ax[idx].set_xlim([minDate, maxDate])
            ax[idx].tick_params(axis='x', rotation=90)
            ax[idx].set_ylim([0, maxContacts])
            ax[idx].set_yticks(yTicks)

        #Save
        figPath = os.path.join(outputPlotFolderPath, fileName + '.png')
        fig.tight_layout()
        fig.savefig(figPath, bbox_inches='tight')

        doc.add_paragraph('The picture below shows {} contacts, depicting the number of satellites in visibility'.format(descr))
        p = doc.add_paragraph()
        r = p.add_run()
        r.add_picture(figPath)

        #Get picutre for average satellite on visibility
        sumNumerosity = np.sum(numerosity)
        if sumNumerosity > 0:
            figV, axV = plt.subplots(figsize=(8, 8))
            for i in range(len(elevs)):
                for j in range(len(lats)):
                    if visibility[i][j]:
                        n = numerosity[i][j] if numerosity[i][j] > 0 else 1
                        visibility[i][j] = visibility[i][j] / n
            pd.DataFrame(visibility, columns=elevs, index=lats).to_csv(os.path.join(outputAnalysFolderPath, 'analysis_' + t + '-sat-visibility-per-lat.csv'))
            cs = axV.contourf(elevs, lats, visibility)
            cf = figV.colorbar(cs, fraction=0.046, pad=0.04)
            cf.ax.set_ylabel("Visible Sats", loc='center')
            axV.set_xlabel("Elevation [deg]")
            axV.set_ylabel("Latitude [deg]")
            axV.set(xlim=[elevs[0], elevs[-1]], ylim=[lats[0], lats[-1]])
            #Save
            fileName = 'analysis_' + t + '-sat-visibility-per-lat'
            figPath = os.path.join(outputPlotFolderPath, fileName + '.png')
            figV.tight_layout()
            figV.savefig(figPath, bbox_inches='tight')
            if sumNumerosity > totNumerosity:
                figPathV = figPath
                totNumerosity = sumNumerosity

    #Save paragraph wtih most numerous sat visibility study
    if figPathV != "":
        doc.add_paragraph('The picture below shows the evolution of number of visible satellites, mean value over the propagation, as varying latitude of the ground point, and min elvation angle.')
        p = doc.add_paragraph()
        r = p.add_run()
        r.add_picture(figPathV)

    plt.close('all')
    print('   - Added section on Ground Stations and User Terminals Contacts in {:.4f} seconds'.format(time.time() - tick))

def getGroundPoint(simulationRequest: dict, poiId: str):
    """ Get from the Simulation Request, the object (GS or UT) based on the id """
    pois = simulationRequest.get('userterminals', []) + simulationRequest.get('groundstations', [])
    for poi in pois:
        if poi['id'] == poiId:
            return poi
    raise Exception('ERROR: ground point with id {} is not defined in Simulation Request, but found in propagation data'.format(poiId))

# -*- coding: utf-8 -*-