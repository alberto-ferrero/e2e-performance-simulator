#! /usr/bin/env python3

# Copyright (C)
# Author: alberto-ferrero

import os
import re
from datetime import datetime

from ...utils.filemanager import getLogoPath, makeOutputFolder, removeFolder

#Import docx
from docx import Document
from docx.shared import Inches, Pt, Mm

#Import analyis
from ..postprocessor.analysis import constellationgeom as constgeom
from ..postprocessor.analysis import gslocation as gsloc
from ..postprocessor.analysis import utlocation as utloc
from ..postprocessor.analysis import contacts as contacts
from ..postprocessor.analysis import links as links
from ..postprocessor.analysis import latency as latency

""" E2E Performance Simulator post processor report writed """

def writeReport(simulationRequest: dict,
                outputDataFolderPath: str,
                outputReportFolderPath: str,
                flightDynamicsDataOutputPath: str,
                regulatoryDataOutputPath: str,
                airLinkDataOutputPath: str,
                spaceLinkDataOutputPath: str,
                networkDataOutputPath: str,
                outputPlotFolderPath: str):
    #Write report based on config request properties
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    section = doc.sections[0]
    section.left_margin = Mm(15)
    section.right_margin = Mm(15)

    #Add logo
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_picture(os.path.join(getLogoPath(), "Logo.png"), width=Inches(3), )

    #First header page
    for i in range(3):
        doc.add_paragraph()
    title = 'E2E Performance Simulator Report'
    currentTime = datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S")
    p = doc.add_heading(title, 0)
    p.bold = True
    for i in range(12):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    r = p.add_run("RSN-SYS")
    r.bold = True
    r.font.size = Pt(20)
    
    p = doc.add_paragraph()
    r = p.add_run(currentTime)
    r.font.size = Pt(11)
    
    doc.add_page_break()

    #Add Section for the E2E Perfomances
    analysis: dict = simulationRequest['analysis']
    for analyisTag in analysis:

        if analyisTag == 'constellation-geometry':
            constgeom.write(doc, simulationRequest, outputPlotFolderPath, flightDynamicsDataOutputPath)

        if analyisTag == 'groundstations-location':
            gsloc.write(doc, outputPlotFolderPath, simulationRequest)

        if analyisTag == 'userterminals-location':
            utloc.write(doc, outputPlotFolderPath, simulationRequest)
  
        if analyisTag == 'contacts':
            contacts.write(doc, outputDataFolderPath, outputPlotFolderPath, flightDynamicsDataOutputPath)
    
        if analyisTag == 'links':
            links.write(doc, outputDataFolderPath, outputPlotFolderPath, airLinkDataOutputPath)
  
        if analyisTag == 'latency':
            latency.write(doc, outputDataFolderPath, outputPlotFolderPath, flightDynamicsDataOutputPath)
  
    #Add Sections for each module
    modules: dict = simulationRequest['modules']
    for moduleTag, module in modules.items():

        if moduleTag == 'flightDynamics' and 'report' in module:
            #Add Flight Dynamics Results
            doc.add_heading("Appendinx: Flight Dynamics", 1)
            p = doc.add_paragraph()
            r = p.add_run("The table blow shows intitial orbit information regarding satellites as defined in the Simulation Request.")

            table = doc.add_table(rows=1, cols=2, style="Table Grid")
            heading = table.rows[0].cells
            heading[0].text = "Satellite id"
            heading[1].text = "Orbit"

            # Get for each satellite orbit plot
            for satellite in simulationRequest.get('satellites', []):
                row = table.add_row().cells
                row[0].text = satellite['id']
                satellite['orbit'].pop('type')
                orbit = row[1].add_table(rows=len(satellite['orbit']), cols=2)
                for i, (k, v) in enumerate(satellite['orbit'].items()):
                    orbit.rows[i].cells[0].text = " ".join(re.sub('([A-Z][a-z]+)', r' \1', re.sub('([A-Z]+)', r' \1', k)).lower().split())
                    orbit.rows[i].cells[1].text = str(v)
            doc.add_page_break()

        if moduleTag == 'regulatorMap' and 'report' in module:
            #Add Regulatory Mapper Results
            #TODO
            pass

        if moduleTag == 'airLinkBudget' and 'report' in module:
            #Add Air Link Budget Results
            #TODO
            pass

        if moduleTag == 'spaceLinkBudget' and 'report' in module:
            #Add Space Link Budget Results
            #TODO
            pass

        if moduleTag == 'networkTopology' and 'report' in module:
            #Add Network Topology Results
            #TODO
            pass

    #Save doc
    makeOutputFolder(outputReportFolderPath)
    doc.save(os.path.join(outputReportFolderPath, "E2E_Performance_Simulator_Report.docx"))
    
# -*- coding: utf-8 -*-